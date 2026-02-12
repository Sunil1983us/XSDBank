#!/usr/bin/env python3
"""
ISO 20022 Payment Schema Comparison Tool - ENHANCED VERSION  
Compares XSD schemas in XML message order with detailed difference reporting
Generates BOTH Excel and Word document reports
"""

import xml.etree.ElementTree as ET
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import argparse
from pathlib import Path
from datetime import datetime


class XSDComparator:
    def _classify_field_from_xsd(self, element, elem_name='', min_occurs='1', annotation=None):
        """Read Yellow/White from XSD annotations - ISO 20022 Spec"""
        # Try to get element if passed as string
        if isinstance(element, str):
            elem_name = element
            element = None
        
        # Check XSD annotation first
        if element is not None:
            ns = getattr(self, 'ns', {'xs': 'http://www.w3.org/2001/XMLSchema'})
            annotation_elem = element.find('xs:annotation', ns)
            if annotation_elem is not None:
                docs = annotation_elem.findall('xs:documentation', ns)
                for doc in docs:
                    source = doc.get('source', '').strip()
                    if source == 'Yellow Field':
                        return '🟡 Yellow (ISO 20022 Spec)'
                    elif source == 'White Field':
                        return '⚪ White (ISO 20022 Spec)'
        
        # NO INFERENCE - Only use XSD annotations
        # If not in XSD, return NA
        
        return '⚫ NA (Not in XSD)'

    """Compare two XSD schemas in XML message order"""
    
    NAMESPACES = {
        'xs': 'http://www.w3.org/2001/XMLSchema',
        'xsd': 'http://www.w3.org/2001/XMLSchema'
    }
    
    def __init__(self, schema1_file, schema2_file, name1=None, name2=None):
        self.schema1_file = schema1_file
        self.schema2_file = schema2_file
        self.name1 = name1 or Path(schema1_file).stem
        self.name2 = name2 or Path(schema2_file).stem
        
        # Parse both schemas
        self.schema1 = self._parse_schema(schema1_file)
        self.schema2 = self._parse_schema(schema2_file)
        
        # Store differences
        self.differences = []
        
        # Build type caches for restriction comparison
        # Get schema roots first
        self.schema1_root = ET.parse(schema1_file).getroot()
        self.schema2_root = ET.parse(schema2_file).getroot()
        self.ns = self.NAMESPACES
        
        self.schema1_type_cache = self._build_type_cache(self.schema1_root, self.ns)
        self.schema2_type_cache = self._build_type_cache(self.schema2_root, self.ns)
        


    def _build_type_cache(self, schema_root, ns):
        """Build cache of all type definitions with their restrictions"""
        type_cache = {}
        
        # Extract simpleType definitions
        for simple_type in schema_root.findall('.//xs:simpleType', ns):
            name = simple_type.get('name')
            if name:
                type_cache[name] = self._extract_type_restrictions(simple_type, ns)
        
        # Extract complexType base restrictions
        for complex_type in schema_root.findall('.//xs:complexType', ns):
            name = complex_type.get('name')
            if name:
                # Check for simpleContent with restriction
                simple_content = complex_type.find('.//xs:simpleContent', ns)
                if simple_content is not None:
                    restriction = simple_content.find('.//xs:restriction', ns)
                    if restriction is not None:
                        type_cache[name] = self._extract_restrictions_from_element(restriction, ns)
        
        return type_cache
    
    def _extract_type_restrictions(self, simple_type, ns):
        """Extract restrictions from a simpleType"""
        restrictions = {}
        restriction = simple_type.find('xs:restriction', ns)
        
        if restriction is not None:
            restrictions = self._extract_restrictions_from_element(restriction, ns)
            restrictions['base'] = restriction.get('base', '').split(':')[-1]
        
        return restrictions
    
    def _extract_restrictions_from_element(self, restriction, ns):
        """Extract all restriction facets"""
        restrictions = {}
        
        # maxLength
        max_len = restriction.find('xs:maxLength', ns)
        if max_len is not None:
            restrictions['maxLength'] = max_len.get('value')
        
        # minLength
        min_len = restriction.find('xs:minLength', ns)
        if min_len is not None:
            restrictions['minLength'] = min_len.get('value')
        
        # length (exact)
        length = restriction.find('xs:length', ns)
        if length is not None:
            restrictions['length'] = length.get('value')
        
        # pattern
        pattern = restriction.find('xs:pattern', ns)
        if pattern is not None:
            restrictions['pattern'] = pattern.get('value')
        
        # minInclusive
        min_inc = restriction.find('xs:minInclusive', ns)
        if min_inc is not None:
            restrictions['minInclusive'] = min_inc.get('value')
        
        # maxInclusive
        max_inc = restriction.find('xs:maxInclusive', ns)
        if max_inc is not None:
            restrictions['maxInclusive'] = max_inc.get('value')
        
        # enumeration
        enums = restriction.findall('xs:enumeration', ns)
        if enums:
            restrictions['enumeration'] = [e.get('value') for e in enums]
        
        # fractionDigits
        frac = restriction.find('xs:fractionDigits', ns)
        if frac is not None:
            restrictions['fractionDigits'] = frac.get('value')
        
        # totalDigits
        total = restriction.find('xs:totalDigits', ns)
        if total is not None:
            restrictions['totalDigits'] = total.get('value')
        
        return restrictions
    
    def _compare_type_restrictions(self, type1, type2):
        """Compare restrictions between two types"""
        rest1 = self.schema1_type_cache.get(type1, {})
        rest2 = self.schema2_type_cache.get(type2, {})
        
        differences = []
        all_keys = set(rest1.keys()) | set(rest2.keys())
        
        for key in all_keys:
            if key == 'enumeration':
                continue  # Skip enum comparison for now
            
            val1 = rest1.get(key)
            val2 = rest2.get(key)
            
            if val1 != val2:
                differences.append(f"{key}: {val1 or 'N/A'} → {val2 or 'N/A'}")
        
        return differences

    def _parse_schema(self, xsd_file):
        """Parse XSD schema and build element structure"""
        tree = ET.parse(xsd_file)
        root = tree.getroot()
        ns_prefix = self._detect_namespace(root)
        
        # Build type cache
        type_cache = {}
        for complex_type in root.findall(f'{ns_prefix}complexType', self.NAMESPACES):
            type_name = complex_type.get('name', '')
            if type_name:
                type_cache[type_name] = complex_type
        
        for simple_type in root.findall(f'{ns_prefix}simpleType', self.NAMESPACES):
            type_name = simple_type.get('name', '')
            if type_name:
                type_cache[type_name] = simple_type
        
        # Extract metadata
        metadata = {
            'target_namespace': root.get('targetNamespace', ''),
            'root': root,
            'ns_prefix': ns_prefix,
            'type_cache': type_cache
        }
        
        # Extract scheme info
        root_elem = root.find(f'{ns_prefix}element', self.NAMESPACES)
        if root_elem is not None:
            root_type = root_elem.get('type', '')
            metadata['root_element'] = root_elem.get('name', '')
            metadata['root_type'] = root_type
            
            # Parse scheme from type name
            if '_' in root_type:
                metadata['scheme'] = root_type.split('_', 1)[-1]
            else:
                metadata['scheme'] = ''
        
        # Parse all elements in XML order
        elements = self._parse_elements(root, ns_prefix, type_cache)
        metadata['elements'] = elements
        
        return metadata
    
    def _detect_namespace(self, root):
        tag = root.tag
        if '{http://www.w3.org/2001/XMLSchema}' in tag:
            return '{http://www.w3.org/2001/XMLSchema}'
        return ''
    
    def _parse_elements(self, root, ns_prefix, type_cache):
        """Parse all elements in XML message order"""
        elements_dict = {}
        sequence = {'count': 0}
        
        # Find root element
        for elem in root.findall(f'{ns_prefix}element', self.NAMESPACES):
            element_name = elem.get('name', '')
            element_type = elem.get('type', '')
            
            self._expand_element(
                elem, element_name, element_type,
                path=element_name, level=0,
                sequence=sequence, elements_dict=elements_dict,
                ns_prefix=ns_prefix, type_cache=type_cache
            )
        
        return elements_dict
    
    def _expand_element(self, element_node, element_name, element_type, path, level, sequence, elements_dict, ns_prefix, type_cache):
        """Expand element following type references"""
        sequence['count'] += 1
        
        # Get element properties
        min_occurs = element_node.get('minOccurs', '1')
        max_occurs = element_node.get('maxOccurs', '1')
        default = element_node.get('default', '')
        fixed = element_node.get('fixed', '')
        
        # Get restrictions
        restrictions = self._get_restrictions(element_node, element_type, ns_prefix, type_cache)
        
        # Get Yellow/White field classification
        field_class = self._classify_field_from_xsd(element_node)
        
        # Get documentation details (Rulebook, Usage Rules)
        documentation = self._extract_documentation(element_node, ns_prefix)
        
        # Get enumeration values if applicable
        enumerations = self._get_enumerations(element_node, element_type, ns_prefix, type_cache)
        
        # Store element info
        elements_dict[path] = {
            'sequence': sequence['count'],
            'name': element_name,
            'path': path,
            'type': element_type,
            'min_occurs': min_occurs,
            'max_occurs': max_occurs,
            'default': default,
            'fixed': fixed,
            'restrictions': restrictions,
            'level': level,
            'node_type': 'element',
            'field_class': field_class,
            'rulebook': documentation.get('rulebook', ''),
            'usage_rules': documentation.get('usage_rules', ''),
            'enumerations': enumerations
        }
        
        # Expand children
        inline_complex = element_node.find(f'{ns_prefix}complexType', self.NAMESPACES)
        if inline_complex is not None:
            self._expand_complex_type(inline_complex, path, level + 1, sequence, elements_dict, ns_prefix, type_cache)
        elif element_type:
            type_def = type_cache.get(element_type)
            if type_def is not None:
                tag_name = self._get_tag_name(type_def)
                if tag_name == 'complexType':
                    self._expand_complex_type(type_def, path, level + 1, sequence, elements_dict, ns_prefix, type_cache)
    
    def _extract_documentation(self, element_node, ns_prefix):
        """Extract Rulebook and Usage Rules from element annotations"""
        result = {'rulebook': '', 'usage_rules': ''}
        
        annotation = element_node.find(f'{ns_prefix}annotation', self.NAMESPACES)
        if annotation is not None:
            usage_rules = []
            for doc in annotation.findall(f'{ns_prefix}documentation', self.NAMESPACES):
                source = doc.get('source', '').strip()
                text = (doc.text or '').strip()
                
                if source == 'Rulebook' and text:
                    result['rulebook'] = text
                elif source == 'Usage Rule' and text:
                    usage_rules.append(text)
                elif source == 'Format Rule' and text:
                    usage_rules.append(f"[Format] {text}")
            
            result['usage_rules'] = ' | '.join(usage_rules) if usage_rules else ''
        
        return result
    
    def _get_enumerations(self, element_node, element_type, ns_prefix, type_cache):
        """Get enumeration values for an element"""
        enums = []
        
        # Check inline simpleType
        simple_type = element_node.find(f'{ns_prefix}simpleType', self.NAMESPACES)
        if simple_type is not None:
            restriction = simple_type.find(f'{ns_prefix}restriction', self.NAMESPACES)
            if restriction is not None:
                for enum in restriction.findall(f'{ns_prefix}enumeration', self.NAMESPACES):
                    enums.append(enum.get('value', ''))
        
        # Check type reference
        if not enums and element_type:
            type_def = type_cache.get(element_type)
            if type_def is not None:
                restriction = type_def.find(f'{ns_prefix}restriction', self.NAMESPACES)
                if restriction is not None:
                    for enum in restriction.findall(f'{ns_prefix}enumeration', self.NAMESPACES):
                        enums.append(enum.get('value', ''))
        
        return sorted(enums) if enums else []
    
    def _expand_complex_type(self, complex_type, parent_path, level, sequence, elements_dict, ns_prefix, type_cache):
        """Expand complex type"""
        # Handle complexContent
        complex_content = complex_type.find(f'{ns_prefix}complexContent', self.NAMESPACES)
        if complex_content is not None:
            restriction = complex_content.find(f'{ns_prefix}restriction', self.NAMESPACES)
            extension = complex_content.find(f'{ns_prefix}extension', self.NAMESPACES)
            
            target = restriction if restriction is not None else extension
            if target is not None:
                base = target.get('base', '')
                base_type_def = type_cache.get(base)
                if base_type_def is not None:
                    self._expand_complex_type(base_type_def, parent_path, level, sequence, elements_dict, ns_prefix, type_cache)
                self._parse_type_content(target, parent_path, level, sequence, elements_dict, ns_prefix, type_cache)
        else:
            self._parse_type_content(complex_type, parent_path, level, sequence, elements_dict, ns_prefix, type_cache)
    
    def _parse_type_content(self, type_node, parent_path, level, sequence, elements_dict, ns_prefix, type_cache):
        """Parse type content"""
        # Handle sequence
        for seq in type_node.findall(f'{ns_prefix}sequence', self.NAMESPACES):
            for child_elem in seq.findall(f'{ns_prefix}element', self.NAMESPACES):
                child_name = child_elem.get('name', '')
                child_type = child_elem.get('type', '')
                child_path = f"{parent_path}/{child_name}"
                
                self._expand_element(
                    child_elem, child_name, child_type, child_path, level,
                    sequence, elements_dict, ns_prefix, type_cache
                )
        
        # Handle choice
        for choice in type_node.findall(f'{ns_prefix}choice', self.NAMESPACES):
            for child_elem in choice.findall(f'{ns_prefix}element', self.NAMESPACES):
                child_name = child_elem.get('name', '')
                child_type = child_elem.get('type', '')
                child_path = f"{parent_path}/{child_name}"
                
                self._expand_element(
                    child_elem, child_name, child_type, child_path, level,
                    sequence, elements_dict, ns_prefix, type_cache
                )
        
        # Handle attributes
        for attr in type_node.findall(f'{ns_prefix}attribute', self.NAMESPACES):
            sequence['count'] += 1
            attr_name = attr.get('name', '')
            attr_type = attr.get('type', '')
            use = attr.get('use', 'optional')
            default = attr.get('default', '')
            fixed = attr.get('fixed', '')
            
            attr_path = f"{parent_path}/@{attr_name}"
            restrictions = self._get_restrictions(attr, attr_type, ns_prefix, type_cache)
            
            elements_dict[attr_path] = {
                'sequence': sequence['count'],
                'name': f"@{attr_name}",
                'path': attr_path,
                'type': attr_type,
                'min_occurs': '1' if use == 'required' else '0',
                'max_occurs': '1',
                'default': default,
                'fixed': fixed,
                'restrictions': restrictions,
                'level': level,
                'node_type': 'attribute'
            }
    
    def _get_restrictions(self, element_node, element_type, ns_prefix, type_cache):
        """Get restrictions for an element"""
        restrictions = []
        
        # Check inline simple type
        simple_type = element_node.find(f'{ns_prefix}simpleType', self.NAMESPACES)
        if simple_type is not None:
            restrictions.extend(self._parse_simple_type(simple_type, ns_prefix))
        elif element_type:
            type_def = type_cache.get(element_type)
            if type_def is not None and self._get_tag_name(type_def) == 'simpleType':
                restrictions.extend(self._parse_simple_type(type_def, ns_prefix))
        
        return ' | '.join(restrictions) if restrictions else ''
    
    def _parse_simple_type(self, simple_type, ns_prefix):
        """Parse simple type restrictions"""
        restrictions = []
        
        restriction = simple_type.find(f'{ns_prefix}restriction', self.NAMESPACES)
        if restriction is not None:
            # Enumerations
            enums = [e.get('value', '') for e in restriction.findall(f'{ns_prefix}enumeration', self.NAMESPACES)]
            if enums:
                restrictions.append(f"Enum: {', '.join(enums[:5])}{'...' if len(enums) > 5 else ''}")
            
            # Pattern
            pattern = restriction.find(f'{ns_prefix}pattern', self.NAMESPACES)
            if pattern is not None:
                restrictions.append(f"Pattern: {pattern.get('value', '')}")
            
            # Length
            for constraint in ['minLength', 'maxLength', 'length']:
                elem = restriction.find(f'{ns_prefix}{constraint}', self.NAMESPACES)
                if elem is not None:
                    restrictions.append(f"{constraint}: {elem.get('value', '')}")
        
        return restrictions
    
    def _get_tag_name(self, element):
        tag = element.tag
        if '}' in tag:
            return tag.split('}')[1]
        return tag
    
    def compare(self):
        """Compare both schemas"""
        elements1 = self.schema1['elements']
        elements2 = self.schema2['elements']
        
        # Get all paths
        all_paths = sorted(set(list(elements1.keys()) + list(elements2.keys())),
                          key=lambda p: (elements1.get(p, {}).get('sequence', 999999),
                                        elements2.get(p, {}).get('sequence', 999999)))
        
        # Compare each path
        for path in all_paths:
            elem1 = elements1.get(path)
            elem2 = elements2.get(path)
            
            if elem1 is None:
                # Only in schema 2
                self.differences.append({
                    'type': 'ADDED',
                    'severity': 'HIGH',
                    'path': path,
                    'element': elem2['name'],
                    'schema1_value': 'NOT PRESENT',
                    'schema2_value': 'PRESENT',
                    'schema1_type': '',
                    'schema2_type': elem2['type'],
                    'schema1_min': '',
                    'schema2_min': elem2['min_occurs'],
                    'schema1_max': '',
                    'schema2_max': elem2['max_occurs'],
                    'impact': f"New field '{path}' added in {self.name2}. May be required in new version.",
                    'sequence1': 0,
                    'sequence2': elem2['sequence']
                })
            elif elem2 is None:
                # Only in schema 1
                self.differences.append({
                    'type': 'REMOVED',
                    'severity': 'HIGH',
                    'path': path,
                    'element': elem1['name'],
                    'schema1_value': 'PRESENT',
                    'schema2_value': 'NOT PRESENT',
                    'schema1_type': elem1['type'],
                    'schema2_type': '',
                    'schema1_min': elem1['min_occurs'],
                    'schema2_min': '',
                    'schema1_max': elem1['max_occurs'],
                    'schema2_max': '',
                    'impact': f"Field '{path}' removed in {self.name2}. Breaking change if field was in use.",
                    'sequence1': elem1['sequence'],
                    'sequence2': 0
                })
            else:
                # Present in both - check differences
                self._compare_elements(elem1, elem2, path)
        
        return self.differences
    
    def _compare_elements(self, elem1, elem2, path):
        """Compare two elements - only report ORDER_CHANGED if there are other changes"""
        
        # Track if we have any substantive changes (not just order)
        has_substantive_change = False
        has_order_change = elem1['sequence'] != elem2['sequence']
        
        # Type change
        if elem1['type'] != elem2['type']:
            has_substantive_change = True
            # Compare type restrictions
            restriction_changes = self._compare_type_restrictions(elem1['type'], elem2['type'])
            restriction_details = '; '.join(restriction_changes) if restriction_changes else ''
            
            impact_msg = f"Data type changed from '{elem1['type']}' to '{elem2['type']}'. May require data conversion."
            if restriction_details:
                impact_msg += f" Restrictions: {restriction_details}"
            
            self.differences.append({
                'type': 'TYPE_CHANGED',
                'severity': 'HIGH',
                'path': path,
                'element': elem1['name'],
                'schema1_value': elem1['type'],
                'schema2_value': elem2['type'],
                'schema1_type': elem1['type'],
                'schema2_type': elem2['type'],
                'schema1_min': elem1['min_occurs'],
                'schema2_min': elem2['min_occurs'],
                'schema1_max': elem1['max_occurs'],
                'schema2_max': elem2['max_occurs'],
                'impact': impact_msg,
                'restriction_details': restriction_details,
                'sequence1': elem1['sequence'],
                'sequence2': elem2['sequence']
            })
        
        # Cardinality change
        if elem1['min_occurs'] != elem2['min_occurs']:
            has_substantive_change = True
            severity = 'HIGH' if (elem1['min_occurs'] == '0' and elem2['min_occurs'] != '0') else 'MEDIUM'
            self.differences.append({
                'type': 'CARDINALITY_CHANGED',
                'severity': severity,
                'path': path,
                'element': elem1['name'],
                'schema1_value': f"min:{elem1['min_occurs']}",
                'schema2_value': f"min:{elem2['min_occurs']}",
                'schema1_type': elem1['type'],
                'schema2_type': elem2['type'],
                'schema1_min': elem1['min_occurs'],
                'schema2_min': elem2['min_occurs'],
                'schema1_max': elem1['max_occurs'],
                'schema2_max': elem2['max_occurs'],
                'impact': f"Field {'is now required' if elem2['min_occurs'] != '0' else 'is now optional'}.",
                'sequence1': elem1['sequence'],
                'sequence2': elem2['sequence']
            })
        
        if elem1['max_occurs'] != elem2['max_occurs']:
            has_substantive_change = True
            self.differences.append({
                'type': 'CARDINALITY_CHANGED',
                'severity': 'MEDIUM',
                'path': path,
                'element': elem1['name'],
                'schema1_value': f"max:{elem1['max_occurs']}",
                'schema2_value': f"max:{elem2['max_occurs']}",
                'schema1_type': elem1['type'],
                'schema2_type': elem2['type'],
                'schema1_min': elem1['min_occurs'],
                'schema2_min': elem2['min_occurs'],
                'schema1_max': elem1['max_occurs'],
                'schema2_max': elem2['max_occurs'],
                'impact': f"Max occurrences changed from {elem1['max_occurs']} to {elem2['max_occurs']}.",
                'sequence1': elem1['sequence'],
                'sequence2': elem2['sequence']
            })
        
        # Restrictions change
        if elem1['restrictions'] != elem2['restrictions']:
            has_substantive_change = True
            self.differences.append({
                'type': 'RESTRICTION_CHANGED',
                'severity': 'HIGH',
                'path': path,
                'element': elem1['name'],
                'schema1_value': elem1['restrictions'] or 'None',
                'schema2_value': elem2['restrictions'] or 'None',
                'schema1_type': elem1['type'],
                'schema2_type': elem2['type'],
                'schema1_min': elem1['min_occurs'],
                'schema2_min': elem2['min_occurs'],
                'schema1_max': elem1['max_occurs'],
                'schema2_max': elem2['max_occurs'],
                'impact': "Validation rules changed. May affect data validation.",
                'sequence1': elem1['sequence'],
                'sequence2': elem2['sequence']
            })
        
        # Yellow/White Field Classification change
        field_class1 = elem1.get('field_class', '⚫ NA (Not in XSD)')
        field_class2 = elem2.get('field_class', '⚫ NA (Not in XSD)')
        
        if field_class1 != field_class2:
            has_substantive_change = True
            # Determine severity based on change type
            if 'Yellow' in field_class1 and 'White' in field_class2:
                severity = 'HIGH'
                impact = f"Field classification DOWNGRADED from Yellow (mandatory for scheme) to White (optional). Business impact: field may no longer be required."
            elif 'White' in field_class1 and 'Yellow' in field_class2:
                severity = 'HIGH'
                impact = f"Field classification UPGRADED from White (optional) to Yellow (mandatory for scheme). Business impact: field is now required."
            elif 'NA' in field_class1 and 'Yellow' in field_class2:
                severity = 'MEDIUM'
                impact = f"Field now classified as Yellow (mandatory for scheme). New ISO 20022 requirement."
            elif 'NA' in field_class1 and 'White' in field_class2:
                severity = 'LOW'
                impact = f"Field now classified as White (optional). New ISO 20022 classification."
            elif 'Yellow' in field_class1 and 'NA' in field_class2:
                severity = 'MEDIUM'
                impact = f"Field classification removed (was Yellow). May indicate field is no longer part of scheme specification."
            elif 'White' in field_class1 and 'NA' in field_class2:
                severity = 'LOW'
                impact = f"Field classification removed (was White). May indicate field is no longer part of scheme specification."
            else:
                severity = 'MEDIUM'
                impact = f"Field classification changed."
            
            self.differences.append({
                'type': 'FIELD_CLASS_CHANGED',
                'severity': severity,
                'path': path,
                'element': elem1['name'],
                'schema1_value': field_class1,
                'schema2_value': field_class2,
                'schema1_type': elem1['type'],
                'schema2_type': elem2['type'],
                'schema1_min': elem1['min_occurs'],
                'schema2_min': elem2['min_occurs'],
                'schema1_max': elem1['max_occurs'],
                'schema2_max': elem2['max_occurs'],
                'impact': impact,
                'sequence1': elem1['sequence'],
                'sequence2': elem2['sequence']
            })
        
        # Fixed value change
        fixed1 = elem1.get('fixed', '')
        fixed2 = elem2.get('fixed', '')
        if fixed1 != fixed2:
            has_substantive_change = True
            self.differences.append({
                'type': 'FIXED_VALUE_CHANGED',
                'severity': 'HIGH',
                'path': path,
                'element': elem1['name'],
                'schema1_value': fixed1 or 'None',
                'schema2_value': fixed2 or 'None',
                'schema1_type': elem1['type'],
                'schema2_type': elem2['type'],
                'schema1_min': elem1['min_occurs'],
                'schema2_min': elem2['min_occurs'],
                'schema1_max': elem1['max_occurs'],
                'schema2_max': elem2['max_occurs'],
                'impact': f"Fixed value changed from '{fixed1 or 'None'}' to '{fixed2 or 'None'}'. Messages must use new fixed value.",
                'sequence1': elem1['sequence'],
                'sequence2': elem2['sequence']
            })
        
        # Default value change
        default1 = elem1.get('default', '')
        default2 = elem2.get('default', '')
        if default1 != default2:
            has_substantive_change = True
            self.differences.append({
                'type': 'DEFAULT_VALUE_CHANGED',
                'severity': 'MEDIUM',
                'path': path,
                'element': elem1['name'],
                'schema1_value': default1 or 'None',
                'schema2_value': default2 or 'None',
                'schema1_type': elem1['type'],
                'schema2_type': elem2['type'],
                'schema1_min': elem1['min_occurs'],
                'schema2_min': elem2['min_occurs'],
                'schema1_max': elem1['max_occurs'],
                'schema2_max': elem2['max_occurs'],
                'impact': f"Default value changed from '{default1 or 'None'}' to '{default2 or 'None'}'.",
                'sequence1': elem1['sequence'],
                'sequence2': elem2['sequence']
            })
        
        # Rulebook note change
        rulebook1 = elem1.get('rulebook', '')
        rulebook2 = elem2.get('rulebook', '')
        if rulebook1 != rulebook2:
            has_substantive_change = True
            if rulebook1 and rulebook2:
                severity = 'MEDIUM'
                impact = f"Rulebook note changed. Review business requirements."
            elif rulebook2 and not rulebook1:
                severity = 'MEDIUM'
                impact = f"New Rulebook note added: {rulebook2[:100]}..."
            else:
                severity = 'LOW'
                impact = f"Rulebook note removed."
            
            self.differences.append({
                'type': 'RULEBOOK_CHANGED',
                'severity': severity,
                'path': path,
                'element': elem1['name'],
                'schema1_value': (rulebook1[:100] + '...') if len(rulebook1) > 100 else rulebook1 or 'None',
                'schema2_value': (rulebook2[:100] + '...') if len(rulebook2) > 100 else rulebook2 or 'None',
                'schema1_type': elem1['type'],
                'schema2_type': elem2['type'],
                'schema1_min': elem1['min_occurs'],
                'schema2_min': elem2['min_occurs'],
                'schema1_max': elem1['max_occurs'],
                'schema2_max': elem2['max_occurs'],
                'impact': impact,
                'sequence1': elem1['sequence'],
                'sequence2': elem2['sequence']
            })
        
        # Usage rules change
        usage1 = elem1.get('usage_rules', '')
        usage2 = elem2.get('usage_rules', '')
        if usage1 != usage2:
            has_substantive_change = True
            if usage1 and usage2:
                severity = 'MEDIUM'
                impact = f"Usage rules changed. Review implementation requirements."
            elif usage2 and not usage1:
                severity = 'MEDIUM'
                impact = f"New usage rules added."
            else:
                severity = 'LOW'
                impact = f"Usage rules removed."
            
            self.differences.append({
                'type': 'USAGE_RULES_CHANGED',
                'severity': severity,
                'path': path,
                'element': elem1['name'],
                'schema1_value': (usage1[:100] + '...') if len(usage1) > 100 else usage1 or 'None',
                'schema2_value': (usage2[:100] + '...') if len(usage2) > 100 else usage2 or 'None',
                'schema1_type': elem1['type'],
                'schema2_type': elem2['type'],
                'schema1_min': elem1['min_occurs'],
                'schema2_min': elem2['min_occurs'],
                'schema1_max': elem1['max_occurs'],
                'schema2_max': elem2['max_occurs'],
                'impact': impact,
                'sequence1': elem1['sequence'],
                'sequence2': elem2['sequence']
            })
        
        # Enumeration change
        enums1 = elem1.get('enumerations', [])
        enums2 = elem2.get('enumerations', [])
        if enums1 != enums2:
            has_substantive_change = True
            added_enums = set(enums2) - set(enums1)
            removed_enums = set(enums1) - set(enums2)
            
            if removed_enums:
                severity = 'HIGH'
                impact = f"Enumeration values removed: {', '.join(sorted(removed_enums))}. Breaking change for existing data."
            elif added_enums:
                severity = 'MEDIUM'
                impact = f"Enumeration values added: {', '.join(sorted(added_enums))}. New values available."
            else:
                severity = 'LOW'
                impact = f"Enumeration values changed."
            
            self.differences.append({
                'type': 'ENUMERATION_CHANGED',
                'severity': severity,
                'path': path,
                'element': elem1['name'],
                'schema1_value': ', '.join(enums1) if enums1 else 'None',
                'schema2_value': ', '.join(enums2) if enums2 else 'None',
                'schema1_type': elem1['type'],
                'schema2_type': elem2['type'],
                'schema1_min': elem1['min_occurs'],
                'schema2_min': elem2['min_occurs'],
                'schema1_max': elem1['max_occurs'],
                'schema2_max': elem2['max_occurs'],
                'impact': impact,
                'sequence1': elem1['sequence'],
                'sequence2': elem2['sequence']
            })
        
        # Only report ORDER_CHANGED if there are other substantive changes
        # (Order changes alone are just noise - they're a side effect of added/removed fields)
        # Note: We don't add ORDER_CHANGED to has_substantive_change since it's not substantive itself


class ComparisonReportGenerator:
    """Generate comprehensive comparison Excel report"""
    
    def __init__(self, comparator, output_file):
        self.comparator = comparator
        self.output_file = output_file
        self.wb = Workbook()
        
    def generate(self):
        """Generate all report sheets"""
        if 'Sheet' in self.wb.sheetnames:
            self.wb.remove(self.wb['Sheet'])
        
        self._create_summary_sheet()
        self._create_detailed_comparison_sheet()
        self._create_side_by_side_sheet()
        self._create_added_fields_sheet()
        self._create_removed_fields_sheet()
        self._create_changed_fields_sheet()
        self._create_type_restriction_sheet()
        self._create_field_classification_sheet()
        self._create_documentation_changes_sheet()
        self._create_enumeration_changes_sheet()
        
        self.wb.save(self.output_file)
        
        print(f"\n✅ Comparison report saved: {self.output_file}")
        print(f"   📊 Total differences: {len(self.comparator.differences)}")
        print(f"   📄 Sheets created: 10")
    
    def _create_summary_sheet(self):
        """Create executive summary"""
        ws = self.wb.create_sheet("Summary", 0)
        
        # Title
        ws['A1'] = "ISO 20022 Payment Schema Comparison - Executive Summary"
        ws['A1'].font = Font(size=16, bold=True, color='FFFFFF')
        ws['A1'].fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
        ws.merge_cells('A1:D1')
        
        # Metadata
        row = 3
        ws[f'A{row}'] = "Report Generated:"
        ws[f'B{row}'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        ws[f'A{row}'].font = Font(bold=True)
        
        row += 1
        ws[f'A{row}'] = "Schema 1:"
        ws[f'B{row}'] = f"{self.comparator.name1} ({self.comparator.schema1.get('scheme', 'N/A')})"
        ws[f'A{row}'].font = Font(bold=True)
        
        row += 1
        ws[f'A{row}'] = "Schema 2:"
        ws[f'B{row}'] = f"{self.comparator.name2} ({self.comparator.schema2.get('scheme', 'N/A')})"
        ws[f'A{row}'].font = Font(bold=True)
        
        # Statistics
        row += 2
        ws[f'A{row}'] = "COMPARISON STATISTICS"
        ws[f'A{row}'].font = Font(size=12, bold=True)
        
        row += 1
        stats = self._calculate_statistics()
        for metric, count in stats.items():
            ws[f'A{row}'] = metric
            ws[f'B{row}'] = count
            row += 1
        
        # Severity breakdown
        row += 1
        ws[f'A{row}'] = "SEVERITY BREAKDOWN"
        ws[f'A{row}'].font = Font(size=12, bold=True)
        
        row += 1
        severity_stats = self._calculate_severity_stats()
        for severity, count in severity_stats.items():
            ws[f'A{row}'] = severity
            ws[f'B{row}'] = count
            
            if severity == 'HIGH':
                ws[f'A{row}'].fill = PatternFill(start_color='FF0000', end_color='FF0000', fill_type='solid')
                ws[f'A{row}'].font = Font(color='FFFFFF', bold=True)
            elif severity == 'MEDIUM':
                ws[f'A{row}'].fill = PatternFill(start_color='FFA500', end_color='FFA500', fill_type='solid')
            elif severity == 'LOW':
                ws[f'A{row}'].fill = PatternFill(start_color='90EE90', end_color='90EE90', fill_type='solid')
            
            row += 1
        
        ws.column_dimensions['A'].width = 30
        ws.column_dimensions['B'].width = 50
    
    def _create_detailed_comparison_sheet(self):
        """Create detailed comparison with all differences"""
        ws = self.wb.create_sheet("All Differences")
        
        headers = ['Severity', 'Change Type', 'Element Path', 'Element', 
                   f'{self.comparator.name1}', f'{self.comparator.name2}', 
                   'Impact', 'Restriction Details', 'Seq1', 'Seq2']
        ws.append(headers)
        
        # Style headers
        for cell in ws[1]:
            cell.font = Font(bold=True, color='FFFFFF')
            cell.fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
            cell.alignment = Alignment(horizontal='center', vertical='center')
        
        # Sort by severity and sequence
        sorted_diffs = sorted(self.comparator.differences, 
                             key=lambda x: (x.get('severity', 'LOW'), 
                                          x.get('sequence1', 0), 
                                          x.get('sequence2', 0)))
        
        for diff in sorted_diffs:
            row = [
                diff.get('severity', ''),
                diff.get('type', ''),
                diff.get('path', ''),
                diff.get('element', ''),
                diff.get('schema1_value', ''),
                diff.get('schema2_value', ''),
                diff.get('impact', ''),
                diff.get('restriction_details', ''),
                diff.get('sequence1', ''),
                diff.get('sequence2', '')
            ]
            ws.append(row)
            
            row_num = ws.max_row
            severity = diff.get('severity', 'LOW')
            
            if severity == 'HIGH':
                ws[f'A{row_num}'].fill = PatternFill(start_color='FF0000', end_color='FF0000', fill_type='solid')
                ws[f'A{row_num}'].font = Font(color='FFFFFF', bold=True)
            elif severity == 'MEDIUM':
                ws[f'A{row_num}'].fill = PatternFill(start_color='FFA500', end_color='FFA500', fill_type='solid')
            elif severity == 'LOW':
                ws[f'A{row_num}'].fill = PatternFill(start_color='90EE90', end_color='90EE90', fill_type='solid')
            
            ws[f'G{row_num}'].alignment = Alignment(wrap_text=True, vertical='top')
            ws[f'H{row_num}'].alignment = Alignment(wrap_text=True, vertical='top')
        
        # Column widths
        ws.column_dimensions['A'].width = 10
        ws.column_dimensions['B'].width = 20
        ws.column_dimensions['C'].width = 50
        ws.column_dimensions['D'].width = 25
        ws.column_dimensions['E'].width = 30
        ws.column_dimensions['F'].width = 30
        ws.column_dimensions['G'].width = 50
        ws.column_dimensions['H'].width = 40  # Restriction Details
        ws.column_dimensions['I'].width = 8
        ws.column_dimensions['J'].width = 8
        
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f"A1:J{ws.max_row}"
    
    def _create_side_by_side_sheet(self):
        """Create side-by-side comparison of all fields"""
        ws = self.wb.create_sheet("Side-by-Side")
        
        headers = ['Path', 'Element', 
                   f'{self.comparator.name1} Type', f'{self.comparator.name1} Min', f'{self.comparator.name1} Max',
                   f'{self.comparator.name2} Type', f'{self.comparator.name2} Min', f'{self.comparator.name2} Max',
                   'Status']
        ws.append(headers)
        
        for cell in ws[1]:
            cell.font = Font(bold=True, color='FFFFFF')
            cell.fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
        
        # Get all paths
        all_paths = sorted(set(list(self.comparator.schema1['elements'].keys()) + 
                              list(self.comparator.schema2['elements'].keys())),
                          key=lambda p: (self.comparator.schema1['elements'].get(p, {}).get('sequence', 999999),
                                        self.comparator.schema2['elements'].get(p, {}).get('sequence', 999999)))
        
        for path in all_paths:
            elem1 = self.comparator.schema1['elements'].get(path)
            elem2 = self.comparator.schema2['elements'].get(path)
            
            if elem1 and elem2:
                status = 'CHANGED' if (elem1['type'] != elem2['type'] or 
                                      elem1['min_occurs'] != elem2['min_occurs'] or
                                      elem1['max_occurs'] != elem2['max_occurs']) else 'SAME'
            elif elem1:
                status = 'REMOVED'
            else:
                status = 'ADDED'
            
            row = [
                path,
                elem1['name'] if elem1 else elem2['name'],
                elem1['type'] if elem1 else '',
                elem1['min_occurs'] if elem1 else '',
                elem1['max_occurs'] if elem1 else '',
                elem2['type'] if elem2 else '',
                elem2['min_occurs'] if elem2 else '',
                elem2['max_occurs'] if elem2 else '',
                status
            ]
            ws.append(row)
            
            row_num = ws.max_row
            if status == 'ADDED':
                ws[f'I{row_num}'].fill = PatternFill(start_color='90EE90', end_color='90EE90', fill_type='solid')
            elif status == 'REMOVED':
                ws[f'I{row_num}'].fill = PatternFill(start_color='FF0000', end_color='FF0000', fill_type='solid')
                ws[f'I{row_num}'].font = Font(color='FFFFFF', bold=True)
            elif status == 'CHANGED':
                ws[f'I{row_num}'].fill = PatternFill(start_color='FFA500', end_color='FFA500', fill_type='solid')
        
        ws.column_dimensions['A'].width = 50
        ws.column_dimensions['B'].width = 25
        for col in ['C', 'D', 'E', 'F', 'G', 'H']:
            ws.column_dimensions[col].width = 20
        ws.column_dimensions['I'].width = 12
        
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f"A1:I{ws.max_row}"
    
    def _create_added_fields_sheet(self):
        """Create sheet with added fields only"""
        ws = self.wb.create_sheet("Added Fields")
        
        headers = ['Path', 'Element', 'Type', 'Min', 'Max', 'Restrictions', 'Impact']
        ws.append(headers)
        
        for cell in ws[1]:
            cell.font = Font(bold=True, color='FFFFFF')
            cell.fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
        
        added = [d for d in self.comparator.differences if d['type'] == 'ADDED']
        for diff in sorted(added, key=lambda x: x.get('sequence2', 0)):
            elem = self.comparator.schema2['elements'].get(diff['path'])
            if elem:
                row = [
                    diff['path'],
                    diff['element'],
                    elem['type'],
                    elem['min_occurs'],
                    elem['max_occurs'],
                    elem.get('restrictions', ''),
                    diff['impact']
                ]
                ws.append(row)
        
        for col in ['A', 'B', 'C', 'D', 'E', 'F', 'G']:
            ws.column_dimensions[col].width = 40
        
        ws.freeze_panes = 'A2'
    
    def _create_removed_fields_sheet(self):
        """Create sheet with removed fields only"""
        ws = self.wb.create_sheet("Removed Fields")
        
        headers = ['Path', 'Element', 'Type', 'Min', 'Max', 'Restrictions', 'Impact']
        ws.append(headers)
        
        for cell in ws[1]:
            cell.font = Font(bold=True, color='FFFFFF')
            cell.fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
        
        removed = [d for d in self.comparator.differences if d['type'] == 'REMOVED']
        for diff in sorted(removed, key=lambda x: x.get('sequence1', 0)):
            elem = self.comparator.schema1['elements'].get(diff['path'])
            if elem:
                row = [
                    diff['path'],
                    diff['element'],
                    elem['type'],
                    elem['min_occurs'],
                    elem['max_occurs'],
                    elem.get('restrictions', ''),
                    diff['impact']
                ]
                ws.append(row)
                
                row_num = ws.max_row
                ws[f'A{row_num}'].fill = PatternFill(start_color='FFE6E6', end_color='FFE6E6', fill_type='solid')
        
        for col in ['A', 'B', 'C', 'D', 'E', 'F', 'G']:
            ws.column_dimensions[col].width = 40
        
        ws.freeze_panes = 'A2'
    
    def _create_changed_fields_sheet(self):
        """Create sheet with changed fields only"""
        ws = self.wb.create_sheet("Changed Fields")
        
        headers = ['Path', 'Element', 'Change Type', 
                   f'{self.comparator.name1}', f'{self.comparator.name2}', 'Impact']
        ws.append(headers)
        
        for cell in ws[1]:
            cell.font = Font(bold=True, color='FFFFFF')
            cell.fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
        
        changed = [d for d in self.comparator.differences 
                  if d['type'] not in ['ADDED', 'REMOVED']]
        for diff in sorted(changed, key=lambda x: x.get('path', '')):
            row = [
                diff['path'],
                diff['element'],
                diff['type'],
                diff['schema1_value'],
                diff['schema2_value'],
                diff['impact']
            ]
            ws.append(row)
        
        for col in ['A', 'B', 'C', 'D', 'E', 'F']:
            ws.column_dimensions[col].width = 40
        
        ws.freeze_panes = 'A2'
    

    def _create_type_restriction_sheet(self):
        """Create detailed type restriction changes sheet"""
        ws = self.wb.create_sheet("Type Restriction Changes")
        
        headers = ['Element', 'Path', f'{self.comparator.name1} Type', 
                   f'{self.comparator.name2} Type', 'Restriction', 
                   'Old Value', 'New Value']
        ws.append(headers)
        
        # Style headers
        for cell in ws[1]:
            cell.font = Font(bold=True, color='FFFFFF')
            cell.fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
            cell.alignment = Alignment(horizontal='center', vertical='center')
        
        # Filter for type changes only
        type_changes = [d for d in self.comparator.differences if d.get('type') == 'TYPE_CHANGED']
        
        for diff in type_changes:
            restriction_details = diff.get('restriction_details', '')
            if restriction_details:
                # Parse restriction details
                restrictions = restriction_details.split('; ')
                for restr in restrictions:
                    if ':' in restr:
                        parts = restr.split(': ')
                        if len(parts) == 2:
                            restriction_name = parts[0]
                            values = parts[1]
                            if ' → ' in values:
                                old_val, new_val = values.split(' → ')
                            else:
                                old_val = values
                                new_val = 'N/A'
                            ws.append([
                                diff.get('element', ''),
                                diff.get('path', ''),
                                diff.get('schema1_type', ''),
                                diff.get('schema2_type', ''),
                                restriction_name,
                                old_val,
                                new_val
                            ])
        
        # Column widths
        ws.column_dimensions['A'].width = 25
        ws.column_dimensions['B'].width = 50
        ws.column_dimensions['C'].width = 30
        ws.column_dimensions['D'].width = 30
        ws.column_dimensions['E'].width = 20
        ws.column_dimensions['F'].width = 20
        ws.column_dimensions['G'].width = 20
        
        ws.freeze_panes = 'A2'
        if ws.max_row > 1:
            ws.auto_filter.ref = f"A1:G{ws.max_row}"
    
    def _create_field_classification_sheet(self):
        """Create Yellow/White field classification changes sheet"""
        ws = self.wb.create_sheet("Field Classification Changes")
        
        headers = ['Severity', 'Element', 'Path', 
                   f'{self.comparator.name1} Classification', 
                   f'{self.comparator.name2} Classification',
                   'Change Direction', 'Business Impact']
        ws.append(headers)
        
        # Style headers
        header_fill = PatternFill(start_color='7030A0', end_color='7030A0', fill_type='solid')  # Purple for classification
        for cell in ws[1]:
            cell.font = Font(bold=True, color='FFFFFF')
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        
        # Filter for field classification changes only
        field_class_changes = [d for d in self.comparator.differences if d.get('type') == 'FIELD_CLASS_CHANGED']
        
        # Sort by severity (HIGH first)
        severity_order = {'HIGH': 0, 'MEDIUM': 1, 'LOW': 2}
        field_class_changes.sort(key=lambda x: severity_order.get(x.get('severity', 'LOW'), 3))
        
        for diff in field_class_changes:
            class1 = diff.get('schema1_value', '')
            class2 = diff.get('schema2_value', '')
            
            # Determine change direction
            if 'Yellow' in class1 and 'White' in class2:
                direction = '🟡→⚪ Downgraded'
            elif 'White' in class1 and 'Yellow' in class2:
                direction = '⚪→🟡 Upgraded'
            elif 'NA' in class1 and 'Yellow' in class2:
                direction = '⚫→🟡 New Yellow'
            elif 'NA' in class1 and 'White' in class2:
                direction = '⚫→⚪ New White'
            elif 'Yellow' in class1 and 'NA' in class2:
                direction = '🟡→⚫ Removed'
            elif 'White' in class1 and 'NA' in class2:
                direction = '⚪→⚫ Removed'
            else:
                direction = 'Changed'
            
            ws.append([
                diff.get('severity', ''),
                diff.get('element', ''),
                diff.get('path', ''),
                class1,
                class2,
                direction,
                diff.get('impact', '')
            ])
            
            # Color code rows by severity
            row_num = ws.max_row
            severity = diff.get('severity', 'LOW')
            if severity == 'HIGH':
                fill_color = 'FFCDD2'  # Light red
            elif severity == 'MEDIUM':
                fill_color = 'FFF9C4'  # Light yellow
            else:
                fill_color = 'C8E6C9'  # Light green
            
            for cell in ws[row_num]:
                cell.fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type='solid')
        
        # Column widths
        ws.column_dimensions['A'].width = 12
        ws.column_dimensions['B'].width = 20
        ws.column_dimensions['C'].width = 55
        ws.column_dimensions['D'].width = 28
        ws.column_dimensions['E'].width = 28
        ws.column_dimensions['F'].width = 18
        ws.column_dimensions['G'].width = 60
        
        ws.freeze_panes = 'A2'
        if ws.max_row > 1:
            ws.auto_filter.ref = f"A1:G{ws.max_row}"
    
    def _create_documentation_changes_sheet(self):
        """Create Rulebook and Usage Rules changes sheet"""
        ws = self.wb.create_sheet("Documentation Changes")
        
        headers = ['Severity', 'Change Type', 'Element', 'Path', 
                   f'{self.comparator.name1}', 
                   f'{self.comparator.name2}',
                   'Business Impact']
        ws.append(headers)
        
        # Style headers
        header_fill = PatternFill(start_color='2E7D32', end_color='2E7D32', fill_type='solid')  # Green for docs
        for cell in ws[1]:
            cell.font = Font(bold=True, color='FFFFFF')
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        
        # Filter for documentation changes
        doc_changes = [d for d in self.comparator.differences 
                       if d.get('type') in ['RULEBOOK_CHANGED', 'USAGE_RULES_CHANGED']]
        
        # Sort by type then severity
        type_order = {'RULEBOOK_CHANGED': 0, 'USAGE_RULES_CHANGED': 1}
        severity_order = {'HIGH': 0, 'MEDIUM': 1, 'LOW': 2}
        doc_changes.sort(key=lambda x: (type_order.get(x.get('type'), 9), 
                                         severity_order.get(x.get('severity', 'LOW'), 3)))
        
        for diff in doc_changes:
            change_type = 'Rulebook' if diff.get('type') == 'RULEBOOK_CHANGED' else 'Usage Rules'
            
            ws.append([
                diff.get('severity', ''),
                change_type,
                diff.get('element', ''),
                diff.get('path', ''),
                diff.get('schema1_value', ''),
                diff.get('schema2_value', ''),
                diff.get('impact', '')
            ])
            
            # Color code by type
            row_num = ws.max_row
            if change_type == 'Rulebook':
                fill_color = 'E8F5E9'  # Light green
            else:
                fill_color = 'FFF3E0'  # Light orange
            
            for cell in ws[row_num]:
                cell.fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type='solid')
                cell.alignment = Alignment(wrap_text=True, vertical='top')
        
        # Column widths
        ws.column_dimensions['A'].width = 10
        ws.column_dimensions['B'].width = 15
        ws.column_dimensions['C'].width = 20
        ws.column_dimensions['D'].width = 50
        ws.column_dimensions['E'].width = 50
        ws.column_dimensions['F'].width = 50
        ws.column_dimensions['G'].width = 50
        
        ws.freeze_panes = 'A2'
        if ws.max_row > 1:
            ws.auto_filter.ref = f"A1:G{ws.max_row}"
    
    def _create_enumeration_changes_sheet(self):
        """Create enumeration changes sheet"""
        ws = self.wb.create_sheet("Enumeration Changes")
        
        headers = ['Severity', 'Element', 'Path', 
                   f'{self.comparator.name1} Values', 
                   f'{self.comparator.name2} Values',
                   'Added Values', 'Removed Values', 'Impact']
        ws.append(headers)
        
        # Style headers
        header_fill = PatternFill(start_color='FF6F00', end_color='FF6F00', fill_type='solid')  # Orange
        for cell in ws[1]:
            cell.font = Font(bold=True, color='FFFFFF')
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        
        # Filter for enumeration changes
        enum_changes = [d for d in self.comparator.differences 
                        if d.get('type') == 'ENUMERATION_CHANGED']
        
        for diff in enum_changes:
            val1 = diff.get('schema1_value', '') or ''
            val2 = diff.get('schema2_value', '') or ''
            
            enums1 = set(val1.split(', ')) if val1 and val1 != 'None' else set()
            enums2 = set(val2.split(', ')) if val2 and val2 != 'None' else set()
            
            added = enums2 - enums1
            removed = enums1 - enums2
            
            ws.append([
                diff.get('severity', ''),
                diff.get('element', ''),
                diff.get('path', ''),
                val1,
                val2,
                ', '.join(sorted(added)) if added else 'None',
                ', '.join(sorted(removed)) if removed else 'None',
                diff.get('impact', '')
            ])
            
            # Color code by severity
            row_num = ws.max_row
            severity = diff.get('severity', 'LOW')
            if severity == 'HIGH':
                fill_color = 'FFCDD2'  # Light red
            elif severity == 'MEDIUM':
                fill_color = 'FFF9C4'  # Light yellow
            else:
                fill_color = 'C8E6C9'  # Light green
            
            for cell in ws[row_num]:
                cell.fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type='solid')
                cell.alignment = Alignment(wrap_text=True, vertical='top')
        
        # Column widths
        ws.column_dimensions['A'].width = 10
        ws.column_dimensions['B'].width = 20
        ws.column_dimensions['C'].width = 45
        ws.column_dimensions['D'].width = 30
        ws.column_dimensions['E'].width = 30
        ws.column_dimensions['F'].width = 25
        ws.column_dimensions['G'].width = 25
        ws.column_dimensions['H'].width = 45
        
        ws.freeze_panes = 'A2'
        if ws.max_row > 1:
            ws.auto_filter.ref = f"A1:H{ws.max_row}"
    
    def _calculate_statistics(self):
        """Calculate statistics"""
        stats = {
            'Total Differences': len(self.comparator.differences),
            'Fields Added': len([d for d in self.comparator.differences if d['type'] == 'ADDED']),
            'Fields Removed': len([d for d in self.comparator.differences if d['type'] == 'REMOVED']),
            'Type Changes': len([d for d in self.comparator.differences if d['type'] == 'TYPE_CHANGED']),
            'Cardinality Changes': len([d for d in self.comparator.differences if d['type'] == 'CARDINALITY_CHANGED']),
            'Restriction Changes': len([d for d in self.comparator.differences if d['type'] == 'RESTRICTION_CHANGED']),
            'Field Class Changes': len([d for d in self.comparator.differences if d['type'] == 'FIELD_CLASS_CHANGED']),
            'Enumeration Changes': len([d for d in self.comparator.differences if d['type'] == 'ENUMERATION_CHANGED']),
            'Rulebook Changes': len([d for d in self.comparator.differences if d['type'] == 'RULEBOOK_CHANGED']),
            'Usage Rule Changes': len([d for d in self.comparator.differences if d['type'] == 'USAGE_RULES_CHANGED']),
            'Fixed Value Changes': len([d for d in self.comparator.differences if d['type'] == 'FIXED_VALUE_CHANGED']),
            'Default Value Changes': len([d for d in self.comparator.differences if d['type'] == 'DEFAULT_VALUE_CHANGED']),
        }
        return stats
    
    def _calculate_severity_stats(self):
        """Calculate severity statistics"""
        stats = {'HIGH': 0, 'MEDIUM': 0, 'LOW': 0}
        for diff in self.comparator.differences:
            severity = diff.get('severity', 'LOW')
            stats[severity] = stats.get(severity, 0) + 1
        return stats




class WordDocumentGenerator:
    """Generate comprehensive Word document report"""
    
    def __init__(self, comparator, output_file):
        self.comparator = comparator
        self.output_file = output_file
        self.doc = Document()
        self._setup_styles()
        
    def _setup_styles(self):
        """Setup custom styles"""
        styles = self.doc.styles
        try:
            title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Arial'
            title_style.font.size = Pt(24)
            title_style.font.bold = True
        except:
            pass  # Style already exists
    
    def generate(self):
        """Generate Word document"""
        self._add_title_page()
        self._add_executive_summary()
        self._add_statistics_section()
        self._add_removed_fields()
        self._add_changed_fields()
        self._add_recommendations()
        
        self.doc.save(self.output_file)
        print(f"\n✅ Word document saved: {self.output_file}")
        print(f"   📄 Sections created: 6")
    
    def _add_title_page(self):
        """Add title page"""
        title = self.doc.add_heading('ISO 20022 Payment Schema Comparison Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        subtitle = self.doc.add_paragraph('Detailed Analysis')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        p1 = self.doc.add_paragraph(f"Schema 1: {self.comparator.name1}")
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        vs = self.doc.add_paragraph('vs')
        vs.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        p2 = self.doc.add_paragraph(f"Schema 2: {self.comparator.name2}")
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        date_p = self.doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d')}")
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_page_break()
    
    def _add_executive_summary(self):
        """Add executive summary"""
        self.doc.add_heading('Executive Summary', 1)
        
        total = len(self.comparator.differences)
        summary = f"This comparison identified {total} differences between the two schemas."
        self.doc.add_paragraph(summary)
        
        stats = self._calculate_statistics()
        self.doc.add_heading('Key Findings', 2)
        for metric, count in stats.items():
            self.doc.add_paragraph(f"{metric}: {count}", style='List Bullet')
    
    def _add_statistics_section(self):
        """Add statistics"""
        self.doc.add_heading('Statistics', 1)
        
        stats = self._calculate_statistics()
        table = self.doc.add_table(rows=len(stats)+1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Count'
        
        for i, (metric, count) in enumerate(stats.items(), 1):
            row = table.rows[i].cells
            row[0].text = metric
            row[1].text = str(count)
    
    def _add_removed_fields(self):
        """Add removed fields"""
        self.doc.add_page_break()
        self.doc.add_heading('Removed Fields', 1)
        
        removed = [d for d in self.comparator.differences if d['type'] == 'REMOVED']
        if removed:
            self.doc.add_paragraph(f"Found {len(removed)} removed fields:")
            for diff in removed[:20]:
                self.doc.add_paragraph(diff['path'], style='List Bullet')
            if len(removed) > 20:
                self.doc.add_paragraph(f"... and {len(removed)-20} more (see Excel report)")
        else:
            self.doc.add_paragraph("No fields were removed.")
    
    def _add_changed_fields(self):
        """Add changed fields"""
        self.doc.add_page_break()
        self.doc.add_heading('Changed Fields', 1)
        
        type_changes = [d for d in self.comparator.differences if d['type'] == 'TYPE_CHANGED']
        if type_changes:
            self.doc.add_heading('Type Changes', 2)
            self.doc.add_paragraph(f"Found {len(type_changes)} type changes:")
            for diff in type_changes:  # Show ALL
                p = self.doc.add_paragraph(style='List Bullet')
                p.add_run(f"{diff['path']}: ").bold = True
                p.add_run(f"{diff['schema1_type']} → {diff['schema2_type']}")
            if len(type_changes) > 15:
                self.doc.add_paragraph(f"... and {len(type_changes)-15} more")
    
    def _add_recommendations(self):
        """Add recommendations"""
        self.doc.add_page_break()
        self.doc.add_heading('Recommendations', 1)
        
        removed_count = len([d for d in self.comparator.differences if d['type'] == 'REMOVED'])
        
        if removed_count > 0:
            self.doc.add_heading('Address Removed Fields', 2)
            self.doc.add_paragraph(f"{removed_count} fields were removed. Actions:")
            self.doc.add_paragraph("• Review usage of removed fields", style='List Bullet')
            self.doc.add_paragraph("• Update message templates", style='List Bullet')
            self.doc.add_paragraph("• Modify validation logic", style='List Bullet')
        
        self.doc.add_heading('Testing', 2)
        self.doc.add_paragraph("• Test with new schema", style='List Bullet')
        self.doc.add_paragraph("• Validate all test cases", style='List Bullet')
        self.doc.add_paragraph("• Check backward compatibility", style='List Bullet')
    
    def _calculate_statistics(self):
        """Calculate statistics"""
        return {
            'Total Differences': len(self.comparator.differences),
            'Fields Added': len([d for d in self.comparator.differences if d['type'] == 'ADDED']),
            'Fields Removed': len([d for d in self.comparator.differences if d['type'] == 'REMOVED']),
            'Type Changes': len([d for d in self.comparator.differences if d['type'] == 'TYPE_CHANGED']),
            'Cardinality Changes': len([d for d in self.comparator.differences if d['type'] == 'CARDINALITY_CHANGED']),
            'Field Class Changes': len([d for d in self.comparator.differences if d['type'] == 'FIELD_CLASS_CHANGED']),
            'Enumeration Changes': len([d for d in self.comparator.differences if d['type'] == 'ENUMERATION_CHANGED']),
            'Rulebook Changes': len([d for d in self.comparator.differences if d['type'] == 'RULEBOOK_CHANGED']),
            'Usage Rule Changes': len([d for d in self.comparator.differences if d['type'] == 'USAGE_RULES_CHANGED']),
        }


def main():
    parser = argparse.ArgumentParser(
        description='Compare XSD schemas - Generates Excel + Word reports'
    )
    parser.add_argument('schema1', help='First XSD file')
    parser.add_argument('schema2', help='Second XSD file')
    parser.add_argument('-o', '--output', help='Output Excel file', 
                       default='xsd_comparison_report.xlsx')
    parser.add_argument('-n1', '--name1', help='Name for first schema')
    parser.add_argument('-n2', '--name2', help='Name for second schema')
    
    args = parser.parse_args()
    
    print(f"\n{'='*70}")
    print("ISO 20022 SCHEMA COMPARISON")
    print("Generates: Excel Report + Word Document")
    print(f"{'='*70}\n")
    
    print(f"📂 Schema 1: {args.schema1}")
    print(f"📂 Schema 2: {args.schema2}")
    print("\n⏳ Parsing schemas...")
    
    # Compare
    comparator = XSDComparator(args.schema1, args.schema2, args.name1, args.name2)
    differences = comparator.compare()
    
    print(f"\n⏳ Generating reports...")
    
    # Generate Excel
    excel_report = ComparisonReportGenerator(comparator, args.output)
    excel_report.generate()
    
    # Generate Word
    word_file = args.output.replace('.xlsx', '.docx')
    word_report = WordDocumentGenerator(comparator, word_file)
    word_report.generate()
    
    # Generate HTML
    html_file = args.output.replace('.xlsx', '.html')
    try:
        import sys
        sys.path.insert(0, str(Path(__file__).parent))
        from html_report_generator import InteractiveHTMLGenerator
        html_gen = InteractiveHTMLGenerator(comparator, html_file)
        html_gen.generate()
    except Exception as e:
        print(f"   ⚠️  HTML generation: {str(e)[:50]}")
    
    print(f"\n{'='*70}")
    print("✅ COMPLETE!")
    print(f"{'='*70}\n")
    print(f"📊 Excel: {args.output}")
    print(f"📄 Word: {word_file}")
    print(f"🌐 HTML: {html_file}")
    print(f"📈 Total differences: {len(differences)}\n")


if __name__ == '__main__':
    main()
