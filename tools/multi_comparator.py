#!/usr/bin/env python3
"""
Multi-Schema Comparison Tool
- Compare 2+ schemas simultaneously  
- Excel + Word + HTML reports for each pairwise comparison
- Master comparison matrix
- Stakeholder-specific reports (Management, Business, Developer, QA)
"""

import sys
import os
from pathlib import Path

# Add tools directory to path to import other modules
sys.path.insert(0, str(Path(__file__).parent))

try:
    from schema_comparator import XSDComparator, ComparisonReportGenerator, WordDocumentGenerator
    from html_report_generator import InteractiveHTMLGenerator
except ImportError as e:
    print(f"⚠️  Warning: Could not import comparison modules: {e}")
    print("   Some features may be limited.")

import xml.etree.ElementTree as ET
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import argparse
from datetime import datetime
from collections import defaultdict


class MultiSchemaComparator:
    """Compare multiple XSD schemas (2+)"""
    
    def __init__(self, schema_files, schema_names=None):
        self.schema_files = schema_files
        self.schema_names = schema_names or [Path(f).stem for f in schema_files]
        self.schemas = []
        self.all_fields = set()
        self.comparison_matrix = {}
        self.pairwise_comparisons = []
        
    def parse_all(self):
        """Parse all schemas"""
        print(f"\n⏳ Parsing {len(self.schema_files)} schemas...")
        
        for i, (file, name) in enumerate(zip(self.schema_files, self.schema_names), 1):
            print(f"   {i}. {name}")
            try:
                # Use XSDComparator to parse
                comp = XSDComparator(file, file, name, name)
                schema_data = comp.schema1
                self.schemas.append({
                    'name': name,
                    'file': file,
                    'elements': schema_data['elements'],
                    'metadata': schema_data
                })
                self.all_fields.update(schema_data['elements'].keys())
            except Exception as e:
                print(f"   ⚠️  Error parsing {name}: {e}")
        
        print(f"   ✅ Found {len(self.all_fields)} unique fields across all schemas")
    
    def build_comparison_matrix(self):
        """Build field-by-field comparison matrix with full metadata"""
        print("\n⏳ Building comparison matrix...")
        
        for field_path in sorted(self.all_fields):
            self.comparison_matrix[field_path] = {}
            
            for schema in self.schemas:
                elem = schema['elements'].get(field_path)
                if elem:
                    self.comparison_matrix[field_path][schema['name']] = {
                        'present': True,
                        'type': elem.get('type', ''),
                        'min_occurs': elem.get('min_occurs', ''),
                        'max_occurs': elem.get('max_occurs', ''),
                        'restrictions': elem.get('restrictions', ''),
                        'field_class': elem.get('field_class', '⚫ NA (Not in XSD)'),
                        'rulebook': elem.get('rulebook', ''),
                        'usage_rules': elem.get('usage_rules', ''),
                        'enumerations': elem.get('enumerations', [])
                    }
                else:
                    self.comparison_matrix[field_path][schema['name']] = {
                        'present': False,
                        'field_class': ''
                    }
        
        print(f"   ✅ Matrix built: {len(self.comparison_matrix)} fields")
    
    def perform_pairwise_comparisons(self):
        """Perform pairwise comparisons between consecutive schemas"""
        print("\n⏳ Performing pairwise comparisons...")
        
        for i in range(len(self.schema_files) - 1):
            schema1_file = self.schema_files[i]
            schema2_file = self.schema_files[i + 1]
            schema1_name = self.schema_names[i]
            schema2_name = self.schema_names[i + 1]
            
            print(f"   Comparing {schema1_name} → {schema2_name}")
            
            try:
                comparator = XSDComparator(schema1_file, schema2_file, schema1_name, schema2_name)
                differences = comparator.compare()
                
                self.pairwise_comparisons.append({
                    'schema1': schema1_name,
                    'schema2': schema2_name,
                    'comparator': comparator,
                    'differences': differences
                })
                
                print(f"      ✅ {len(differences)} differences found")
            except Exception as e:
                print(f"      ⚠️  Error: {e}")
        
        print(f"   ✅ Completed {len(self.pairwise_comparisons)} pairwise comparisons")


class EnhancedReportGenerator:
    """Generate comprehensive reports for multi-schema comparison"""
    
    def __init__(self, multi_comparator, output_base):
        self.multi_comparator = multi_comparator
        self.output_base = output_base
        self.generated_files = []
        
    def generate_all_reports(self):
        """Generate all reports"""
        print("\n⏳ Generating comprehensive reports...")
        
        # 1. Generate pairwise comparison reports (Excel + Word + HTML)
        for i, comparison in enumerate(self.multi_comparator.pairwise_comparisons):
            # Sanitize names for filenames
            name1 = self._sanitize_filename(comparison['schema1'])
            name2 = self._sanitize_filename(comparison['schema2'])
            output_file = f"{self.output_base}_{name1}_vs_{name2}.xlsx"
            self._generate_pairwise_report(comparison, output_file)
        
        # 2. Generate master matrix (Excel)
        self._generate_master_matrix()
        
        # 3. Generate stakeholder reports (Word)
        self._generate_stakeholder_reports()
        
        print("\n✅ All reports generated!")
        return self.generated_files
    
    def _sanitize_filename(self, name):
        """Sanitize string for use in filename"""
        # Replace spaces and special chars with underscores
        import re
        return re.sub(r'[^\w\-]', '_', name)
    
    def _generate_pairwise_report(self, comparison, output_file):
        """Generate full comparison report (Excel + Word + HTML)"""
        comparator = comparison['comparator']
        
        try:
            # Excel report
            excel_gen = ComparisonReportGenerator(comparator, output_file)
            excel_gen.generate()
            self.generated_files.append(output_file)
            
            # Word document
            word_file = output_file.replace('.xlsx', '.docx')
            word_gen = WordDocumentGenerator(comparator, word_file)
            word_gen.generate()
            self.generated_files.append(word_file)
            
            # HTML report
            html_file = output_file.replace('.xlsx', '.html')
            try:
                html_gen = InteractiveHTMLGenerator(comparator, html_file)
                html_gen.generate()
                self.generated_files.append(html_file)
            except Exception as e:
                print(f"   ⚠️  HTML generation skipped: {e}")
            
        except Exception as e:
            print(f"   ⚠️  Error generating report: {e}")
    
    def _generate_master_matrix(self):
        """Generate comprehensive master comparison matrix with ALL differences"""
        filename = f"{self.output_base}_MASTER_MATRIX.xlsx"
        wb = Workbook()
        
        # Build a lookup of all differences by path for quick access
        differences_by_path = defaultdict(list)
        for comparison in self.multi_comparator.pairwise_comparisons:
            comp_name = f"{comparison['schema1']} → {comparison['schema2']}"
            for diff in comparison['differences']:
                differences_by_path[diff['path']].append({
                    'comparison': comp_name,
                    'type': diff['type'],
                    'severity': diff.get('severity', 'LOW'),
                    'schema1_value': diff.get('schema1_value', ''),
                    'schema2_value': diff.get('schema2_value', ''),
                    'impact': diff.get('impact', '')
                })
        
        # ===== SHEET 1: Full Comparison Matrix =====
        ws = wb.active
        ws.title = "Full Comparison"
        
        # Headers with all details per schema
        headers = ['Field Path', 'Element']
        for schema in self.multi_comparator.schemas:
            headers.extend([
                f"{schema['name']} Present",
                f"{schema['name']} Type",
                f"{schema['name']} Min",
                f"{schema['name']} Max",
                f"{schema['name']} Class"
            ])
        headers.extend(['Change Types', 'Severity', 'Change Summary'])
        ws.append(headers)
        
        # Style headers
        header_fill = PatternFill(start_color='1F4E78', end_color='1F4E78', fill_type='solid')
        for cell in ws[1]:
            cell.font = Font(bold=True, color='FFFFFF', size=9)
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center', wrap_text=True)
        
        # Data rows
        for field_path in sorted(self.multi_comparator.comparison_matrix.keys()):
            schema_data = self.multi_comparator.comparison_matrix[field_path]
            element_name = field_path.split('/')[-1]
            row = [field_path, element_name]
            
            for schema in self.multi_comparator.schemas:
                data = schema_data[schema['name']]
                if data['present']:
                    row.extend([
                        '✓',
                        data.get('type', '')[:35],
                        data.get('min_occurs', ''),
                        data.get('max_occurs', ''),
                        data.get('field_class', '')[:20] if data.get('field_class') else ''
                    ])
                else:
                    row.extend(['✗', '', '', '', ''])
            
            # Get differences for this path
            diffs = differences_by_path.get(field_path, [])
            if diffs:
                change_types = list(set(d['type'] for d in diffs))
                severities = list(set(d['severity'] for d in diffs))
                max_severity = 'HIGH' if 'HIGH' in severities else ('MEDIUM' if 'MEDIUM' in severities else 'LOW')
                
                # Build change summary
                summaries = []
                for d in diffs[:3]:  # Limit to 3 changes in summary
                    summaries.append(f"{d['type']}: {d['schema1_value'][:20] if d['schema1_value'] else 'N/A'} → {d['schema2_value'][:20] if d['schema2_value'] else 'N/A'}")
                if len(diffs) > 3:
                    summaries.append(f"...+{len(diffs)-3} more")
                
                row.extend([
                    ', '.join(change_types),
                    max_severity,
                    ' | '.join(summaries)
                ])
            else:
                row.extend(['No Change', '', ''])
            
            ws.append(row)
            
            # Color coding based on severity
            if diffs:
                max_sev = row[-2]
                if max_sev == 'HIGH':
                    fill_color = 'FFCDD2'  # Light red
                elif max_sev == 'MEDIUM':
                    fill_color = 'FFF9C4'  # Light yellow
                else:
                    fill_color = 'E8F5E9'  # Light green
                for cell in ws[ws.max_row]:
                    cell.fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type='solid')
        
        # Column widths
        ws.column_dimensions['A'].width = 45
        ws.column_dimensions['B'].width = 18
        col_idx = 3
        for _ in self.multi_comparator.schemas:
            ws.column_dimensions[get_column_letter(col_idx)].width = 8      # Present
            ws.column_dimensions[get_column_letter(col_idx + 1)].width = 28  # Type
            ws.column_dimensions[get_column_letter(col_idx + 2)].width = 6   # Min
            ws.column_dimensions[get_column_letter(col_idx + 3)].width = 6   # Max
            ws.column_dimensions[get_column_letter(col_idx + 4)].width = 18  # Class
            col_idx += 5
        ws.column_dimensions[get_column_letter(col_idx)].width = 25      # Change Types
        ws.column_dimensions[get_column_letter(col_idx + 1)].width = 10  # Severity
        ws.column_dimensions[get_column_letter(col_idx + 2)].width = 50  # Change Summary
        
        ws.freeze_panes = 'C2'
        ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{ws.max_row}"
        
        # ===== SHEET 2: All Changes Detail =====
        ws_changes = wb.create_sheet("All Changes")
        
        change_headers = ['Field Path', 'Element', 'Comparison', 'Change Type', 'Severity',
                          'Old Value', 'New Value', 'Impact']
        ws_changes.append(change_headers)
        
        for cell in ws_changes[1]:
            cell.font = Font(bold=True, color='FFFFFF')
            cell.fill = PatternFill(start_color='C62828', end_color='C62828', fill_type='solid')
        
        # Add all differences
        for field_path, diffs in sorted(differences_by_path.items()):
            element_name = field_path.split('/')[-1]
            for diff in diffs:
                ws_changes.append([
                    field_path,
                    element_name,
                    diff['comparison'],
                    diff['type'],
                    diff['severity'],
                    str(diff['schema1_value'])[:50] if diff['schema1_value'] else '',
                    str(diff['schema2_value'])[:50] if diff['schema2_value'] else '',
                    diff['impact'][:80] if diff['impact'] else ''
                ])
                
                # Color by severity
                sev = diff['severity']
                if sev == 'HIGH':
                    fill = PatternFill(start_color='FFCDD2', end_color='FFCDD2', fill_type='solid')
                elif sev == 'MEDIUM':
                    fill = PatternFill(start_color='FFF9C4', end_color='FFF9C4', fill_type='solid')
                else:
                    fill = PatternFill(start_color='E8F5E9', end_color='E8F5E9', fill_type='solid')
                for cell in ws_changes[ws_changes.max_row]:
                    cell.fill = fill
        
        ws_changes.column_dimensions['A'].width = 45
        ws_changes.column_dimensions['B'].width = 18
        ws_changes.column_dimensions['C'].width = 30
        ws_changes.column_dimensions['D'].width = 22
        ws_changes.column_dimensions['E'].width = 10
        ws_changes.column_dimensions['F'].width = 30
        ws_changes.column_dimensions['G'].width = 30
        ws_changes.column_dimensions['H'].width = 50
        
        ws_changes.freeze_panes = 'A2'
        if ws_changes.max_row > 1:
            ws_changes.auto_filter.ref = f"A1:H{ws_changes.max_row}"
        
        # ===== SHEET 3: Changes by Type =====
        ws_by_type = wb.create_sheet("Changes by Type")
        
        # Group changes by type
        changes_by_type = defaultdict(list)
        for field_path, diffs in differences_by_path.items():
            for diff in diffs:
                changes_by_type[diff['type']].append({
                    'path': field_path,
                    'element': field_path.split('/')[-1],
                    **diff
                })
        
        type_headers = ['Change Type', 'Count', 'HIGH', 'MEDIUM', 'LOW', 'Sample Fields']
        ws_by_type.append(type_headers)
        
        for cell in ws_by_type[1]:
            cell.font = Font(bold=True, color='FFFFFF')
            cell.fill = PatternFill(start_color='2E7D32', end_color='2E7D32', fill_type='solid')
        
        for change_type in sorted(changes_by_type.keys()):
            items = changes_by_type[change_type]
            high = len([i for i in items if i['severity'] == 'HIGH'])
            med = len([i for i in items if i['severity'] == 'MEDIUM'])
            low = len([i for i in items if i['severity'] == 'LOW'])
            samples = ', '.join([i['element'] for i in items[:5]])
            if len(items) > 5:
                samples += f' (+{len(items)-5} more)'
            
            ws_by_type.append([change_type, len(items), high, med, low, samples])
        
        ws_by_type.column_dimensions['A'].width = 25
        ws_by_type.column_dimensions['B'].width = 10
        ws_by_type.column_dimensions['C'].width = 8
        ws_by_type.column_dimensions['D'].width = 10
        ws_by_type.column_dimensions['E'].width = 8
        ws_by_type.column_dimensions['F'].width = 60
        
        # ===== SHEET 4: Field Classifications =====
        ws_class = wb.create_sheet("Field Classifications")
        
        headers_class = ['Field Path', 'Element'] + [s['name'] for s in self.multi_comparator.schemas] + ['Classification Changed', 'Evolution']
        ws_class.append(headers_class)
        
        for cell in ws_class[1]:
            cell.font = Font(bold=True, color='FFFFFF')
            cell.fill = PatternFill(start_color='7030A0', end_color='7030A0', fill_type='solid')
        
        for field_path, schema_data in self.multi_comparator.comparison_matrix.items():
            has_classification = False
            classes = []
            for schema in self.multi_comparator.schemas:
                fc = schema_data[schema['name']].get('field_class', '')
                if fc and ('Yellow' in fc or 'White' in fc):
                    has_classification = True
                classes.append(fc if fc else 'N/A')
            
            if has_classification:
                element_name = field_path.split('/')[-1]
                unique = list(dict.fromkeys([c for c in classes if c and c != 'N/A']))
                
                if len(unique) > 1:
                    class_changed = '⚠️ YES'
                    evolution = " → ".join([c.split()[0] if c and c != 'N/A' else '?' for c in classes])
                else:
                    class_changed = 'No'
                    evolution = 'Stable'
                
                row = [field_path, element_name] + classes + [class_changed, evolution]
                ws_class.append(row)
                
                if class_changed == '⚠️ YES':
                    for cell in ws_class[ws_class.max_row]:
                        cell.fill = PatternFill(start_color='FCE4D6', end_color='FCE4D6', fill_type='solid')
        
        ws_class.column_dimensions['A'].width = 45
        ws_class.column_dimensions['B'].width = 18
        for i in range(3, len(headers_class) + 1):
            ws_class.column_dimensions[get_column_letter(i)].width = 25
        
        ws_class.freeze_panes = 'C2'
        if ws_class.max_row > 1:
            ws_class.auto_filter.ref = f"A1:{get_column_letter(len(headers_class))}{ws_class.max_row}"
        
        # ===== SHEET 5: Summary =====
        ws_summary = wb.create_sheet("Summary")
        
        ws_summary.append(['MULTI-SCHEMA COMPARISON SUMMARY', ''])
        ws_summary['A1'].font = Font(bold=True, size=14)
        ws_summary.append([])
        
        ws_summary.append(['General Statistics', ''])
        ws_summary['A3'].font = Font(bold=True)
        ws_summary.append(['Total Unique Fields', len(self.multi_comparator.comparison_matrix)])
        ws_summary.append(['Schemas Compared', len(self.multi_comparator.schemas)])
        ws_summary.append(['Pairwise Comparisons', len(self.multi_comparator.pairwise_comparisons)])
        ws_summary.append(['Total Differences Found', sum(len(c['differences']) for c in self.multi_comparator.pairwise_comparisons)])
        ws_summary.append([])
        
        # Changes by type summary
        ws_summary.append(['Changes by Type', ''])
        ws_summary[f'A{ws_summary.max_row}'].font = Font(bold=True)
        for change_type, items in sorted(changes_by_type.items(), key=lambda x: -len(x[1])):
            ws_summary.append([f"  {change_type}", len(items)])
        ws_summary.append([])
        
        # Classification statistics per schema
        ws_summary.append(['Field Classifications by Schema', ''])
        ws_summary[f'A{ws_summary.max_row}'].font = Font(bold=True)
        for schema in self.multi_comparator.schemas:
            yellow_count = 0
            white_count = 0
            for field_path, schema_data in self.multi_comparator.comparison_matrix.items():
                fc = schema_data[schema['name']].get('field_class', '')
                if 'Yellow' in fc:
                    yellow_count += 1
                elif 'White' in fc:
                    white_count += 1
            ws_summary.append([f"  {schema['name']}", f"🟡 {yellow_count} Yellow, ⚪ {white_count} White"])
        
        ws_summary.column_dimensions['A'].width = 40
        ws_summary.column_dimensions['B'].width = 35
        
        wb.save(filename)
        self.generated_files.append(filename)
        print(f"   ✅ Master matrix: {filename}")
    
    def _generate_stakeholder_reports(self):
        """Generate stakeholder-specific reports"""
        
        # Management Summary
        self._generate_management_summary()
        
        # Business Impact
        self._generate_business_impact()
        
        # Developer Guide  
        self._generate_developer_guide()
        
        # QA Checklist
        self._generate_qa_checklist()
    
    def _generate_management_summary(self):
        """Generate management summary with field classification insights"""
        filename = f"{self.output_base}_MANAGEMENT_SUMMARY.docx"
        doc = Document()
        
        # Title
        title = doc.add_heading('Multi-Schema Comparison Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Executive Summary")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph()
        
        # Schemas compared
        doc.add_heading('Schemas Compared', 1)
        for i, schema in enumerate(self.multi_comparator.schemas, 1):
            doc.add_paragraph(f"{i}. {schema['name']}", style='List Number')
        
        # Key metrics
        doc.add_paragraph()
        doc.add_heading('Key Metrics', 1)
        doc.add_paragraph(f"• Total unique fields analyzed: {len(self.multi_comparator.all_fields)}")
        doc.add_paragraph(f"• Pairwise comparisons performed: {len(self.multi_comparator.pairwise_comparisons)}")
        
        # Field Classification Statistics
        doc.add_paragraph()
        doc.add_heading('Field Classification Statistics (Yellow/White)', 1)
        
        for schema in self.multi_comparator.schemas:
            yellow_count = 0
            white_count = 0
            for field_path, schema_data in self.multi_comparator.comparison_matrix.items():
                fc = schema_data[schema['name']].get('field_class', '')
                if 'Yellow' in fc:
                    yellow_count += 1
                elif 'White' in fc:
                    white_count += 1
            
            doc.add_paragraph(f"{schema['name']}:", style='List Bullet')
            doc.add_paragraph(f"  🟡 Yellow (Mandatory): {yellow_count} fields")
            doc.add_paragraph(f"  ⚪ White (Optional): {white_count} fields")
        
        # Classification changes across versions
        class_changes = 0
        for field_path, schema_data in self.multi_comparator.comparison_matrix.items():
            classes = set()
            for schema in self.multi_comparator.schemas:
                fc = schema_data[schema['name']].get('field_class', '')
                if fc and 'NA' not in fc:
                    classes.add(fc)
            if len(classes) > 1:
                class_changes += 1
        
        doc.add_paragraph()
        doc.add_paragraph(f"⚠️ Fields with classification changes across versions: {class_changes}", 
                         style='Intense Quote')
        
        # Comparison summary
        doc.add_paragraph()
        doc.add_heading('Comparison Summary', 1)
        
        total_differences = 0
        for comparison in self.multi_comparator.pairwise_comparisons:
            diff_count = len(comparison['differences'])
            total_differences += diff_count
            
            # Count by type
            added = len([d for d in comparison['differences'] if d['type'] == 'ADDED'])
            removed = len([d for d in comparison['differences'] if d['type'] == 'REMOVED'])
            field_class = len([d for d in comparison['differences'] if d['type'] == 'FIELD_CLASS_CHANGED'])
            
            doc.add_paragraph(
                f"{comparison['schema1']} → {comparison['schema2']}: "
                f"{diff_count} differences (Added: {added}, Removed: {removed}, Classification: {field_class})",
                style='List Bullet'
            )
        
        doc.add_paragraph()
        doc.add_paragraph(f"Total differences across all comparisons: {total_differences}", 
                         style='Intense Quote')
        
        doc.save(filename)
        self.generated_files.append(filename)
        print(f"   ✅ Management summary: {filename}")
    
    def _generate_business_impact(self):
        """Generate business impact report with field classification analysis"""
        filename = f"{self.output_base}_BUSINESS_IMPACT.docx"
        doc = Document()
        
        doc.add_heading('Business Impact Report', 0)
        doc.add_paragraph("Analysis of changes and their business implications")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        
        for comparison in self.multi_comparator.pairwise_comparisons:
            doc.add_page_break()
            doc.add_heading(f"{comparison['schema1']} → {comparison['schema2']}", 1)
            
            differences = comparison['differences']
            
            # Categorize by severity
            high_severity = [d for d in differences if d.get('severity') == 'HIGH']
            medium_severity = [d for d in differences if d.get('severity') == 'MEDIUM']
            low_severity = [d for d in differences if d.get('severity') == 'LOW']
            
            # Field classification changes (important for business)
            field_class_changes = [d for d in differences if d.get('type') == 'FIELD_CLASS_CHANGED']
            yellow_to_white = [d for d in field_class_changes if 'Yellow' in str(d.get('schema1_value', '')) and 'White' in str(d.get('schema2_value', ''))]
            white_to_yellow = [d for d in field_class_changes if 'White' in str(d.get('schema1_value', '')) and 'Yellow' in str(d.get('schema2_value', ''))]
            new_yellow = [d for d in field_class_changes if 'NA' in str(d.get('schema1_value', '')) and 'Yellow' in str(d.get('schema2_value', ''))]
            
            doc.add_paragraph(f"Total changes: {len(differences)}")
            doc.add_paragraph(f"  • High: {len(high_severity)}, Medium: {len(medium_severity)}, Low: {len(low_severity)}")
            
            # Field Classification Impact Section
            if field_class_changes:
                doc.add_heading('Field Classification Impact', 2)
                doc.add_paragraph(f"Found {len(field_class_changes)} classification changes:")
                
                if white_to_yellow:
                    doc.add_paragraph(f"⚠️ UPGRADED to Yellow (Now Mandatory): {len(white_to_yellow)} fields", style='List Bullet')
                    for diff in white_to_yellow[:10]:
                        doc.add_paragraph(f"    • {diff['path']}")
                
                if yellow_to_white:
                    doc.add_paragraph(f"ℹ️ DOWNGRADED to White (Now Optional): {len(yellow_to_white)} fields", style='List Bullet')
                    for diff in yellow_to_white[:10]:
                        doc.add_paragraph(f"    • {diff['path']}")
                
                if new_yellow:
                    doc.add_paragraph(f"🆕 NEW Yellow Fields (New Mandatory): {len(new_yellow)} fields", style='List Bullet')
                    for diff in new_yellow[:10]:
                        doc.add_paragraph(f"    • {diff['path']}")
            
            if high_severity:
                doc.add_heading('High Priority Changes', 2)
                doc.add_paragraph(f"Found {len(high_severity)} critical changes requiring immediate attention:")
                
                # Group by type
                by_type = defaultdict(list)
                for diff in high_severity:
                    by_type[diff['type']].append(diff)
                
                for diff_type, diffs in sorted(by_type.items()):
                    doc.add_paragraph(f"{diff_type}: {len(diffs)} changes", style='List Bullet')
                    for diff in diffs[:5]:
                        doc.add_paragraph(f"    • {diff['path']}")
                    if len(diffs) > 5:
                        doc.add_paragraph(f"    ... and {len(diffs) - 5} more")
            
            if medium_severity:
                doc.add_heading('Medium Priority Changes', 2)
                doc.add_paragraph(f"Found {len(medium_severity)} changes requiring review:")
                
                by_type = defaultdict(list)
                for diff in medium_severity:
                    by_type[diff['type']].append(diff)
                
                for diff_type, diffs in sorted(by_type.items()):
                    doc.add_paragraph(f"{diff_type}: {len(diffs)} changes", style='List Bullet')
            
            if not high_severity and not medium_severity:
                doc.add_paragraph("No high or medium priority changes found.")
        
        doc.save(filename)
        self.generated_files.append(filename)
        print(f"   ✅ Business impact: {filename}")
    
    def _generate_developer_guide(self):
        """Generate developer guide"""
        filename = f"{self.output_base}_DEVELOPER_GUIDE.docx"
        doc = Document()
        
        doc.add_heading('Developer Migration Guide', 0)
        doc.add_paragraph("Technical implementation details for schema transitions")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        
        for comparison in self.multi_comparator.pairwise_comparisons:
            doc.add_page_break()
            doc.add_heading(f"Migrating: {comparison['schema1']} → {comparison['schema2']}", 1)
            
            differences = comparison['differences']
            removed = [d for d in differences if d['type'] == 'REMOVED']
            added = [d for d in differences if d['type'] == 'ADDED']
            modified = [d for d in differences if d['type'] not in ['REMOVED', 'ADDED']]
            
            if removed:
                doc.add_heading('Fields Removed (Action: Delete)', 2)
                for diff in removed[:30]:
                    doc.add_paragraph(f"• {diff['path']}", style='List Bullet')
                if len(removed) > 30:
                    doc.add_paragraph(f"... and {len(removed) - 30} more fields")
            
            if added:
                doc.add_heading('Fields Added (Action: Implement)', 2)
                for diff in added[:30]:
                    doc.add_paragraph(f"• {diff['path']}", style='List Bullet')
                if len(added) > 30:
                    doc.add_paragraph(f"... and {len(added) - 30} more fields")
            
            if modified:
                doc.add_heading('Fields Modified (Action: Update)', 2)
                for diff in modified[:30]:
                    doc.add_paragraph(f"• {diff['path']}: {diff['type']}", style='List Bullet')
                if len(modified) > 30:
                    doc.add_paragraph(f"... and {len(modified) - 30} more fields")
            
            if not differences:
                doc.add_paragraph("No differences found - schemas are identical.")
        
        doc.save(filename)
        self.generated_files.append(filename)
        print(f"   ✅ Developer guide: {filename}")
    
    def _generate_qa_checklist(self):
        """Generate QA checklist"""
        filename = f"{self.output_base}_QA_CHECKLIST.docx"
        doc = Document()
        
        doc.add_heading('QA Testing Checklist', 0)
        doc.add_paragraph("Comprehensive testing requirements for schema migrations")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        
        for comparison in self.multi_comparator.pairwise_comparisons:
            doc.add_page_break()
            doc.add_heading(f"Testing: {comparison['schema1']} → {comparison['schema2']}", 1)
            
            differences = comparison['differences']
            
            doc.add_heading('Pre-Migration Testing', 2)
            scenarios = [
                "☐ Backup existing production data",
                "☐ Document current system behavior",
                "☐ Identify all affected integrations",
            ]
            for scenario in scenarios:
                doc.add_paragraph(scenario, style='List Bullet')
            
            doc.add_heading('Schema Validation Testing', 2)
            scenarios = [
                "☐ Validate sample messages against new schema",
                f"☐ Test {len([d for d in differences if d['type'] == 'ADDED'])} new fields with valid data",
                f"☐ Verify {len([d for d in differences if d['type'] == 'REMOVED'])} removed fields cause validation errors",
                "☐ Test boundary conditions for modified fields",
            ]
            for scenario in scenarios:
                doc.add_paragraph(scenario, style='List Bullet')
            
            doc.add_heading('Integration Testing', 2)
            scenarios = [
                "☐ Test upstream system compatibility",
                "☐ Test downstream system compatibility",
                "☐ Verify error handling for rejected messages",
                "☐ End-to-end transaction testing",
            ]
            for scenario in scenarios:
                doc.add_paragraph(scenario, style='List Bullet')
            
            doc.add_heading('Post-Migration Verification', 2)
            scenarios = [
                "☐ Monitor error rates for 24-48 hours",
                "☐ Verify all business functions operational",
                "☐ Confirm reporting accuracy",
                "☐ Document any issues encountered",
            ]
            for scenario in scenarios:
                doc.add_paragraph(scenario, style='List Bullet')
        
        doc.save(filename)
        self.generated_files.append(filename)
        print(f"   ✅ QA checklist: {filename}")


def main():
    parser = argparse.ArgumentParser(
        description='Multi-Schema Comparison Tool - Compare 2+ XSD schemas',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Compare 2 schemas (full reports)
  python %(prog)s schema1.xsd schema2.xsd -o comparison
  
  # Compare 4 versions with custom names
  python %(prog)s v1.xsd v2.xsd v3.xsd v4.xsd -n "V1.0" "V2.0" "V3.0" "V4.0" -o evolution
  
  # Compare vendor schemas
  python %(prog)s vendor_a.xsd vendor_b.xsd vendor_c.xsd -o vendors
        """
    )
    
    parser.add_argument('schemas', nargs='+', help='XSD files to compare (2 or more)')
    parser.add_argument('-o', '--output', help='Output base name', default='comparison')
    parser.add_argument('-n', '--names', nargs='+', help='Names for schemas (optional)')
    
    args = parser.parse_args()
    
    if len(args.schemas) < 2:
        print("❌ Error: Need at least 2 schemas to compare")
        sys.exit(1)
    
    # Validate files exist
    for schema_file in args.schemas:
        if not Path(schema_file).exists():
            print(f"❌ Error: File not found: {schema_file}")
            sys.exit(1)
    
    print(f"\n{'='*70}")
    print("MULTI-SCHEMA COMPARISON TOOL")
    print(f"Comparing {len(args.schemas)} schemas")
    print("Generates: Excel + Word + HTML reports")
    print(f"{'='*70}")
    
    # Parse all schemas
    comparator = MultiSchemaComparator(args.schemas, args.names)
    comparator.parse_all()
    comparator.build_comparison_matrix()
    comparator.perform_pairwise_comparisons()
    
    # Generate reports
    reporter = EnhancedReportGenerator(comparator, args.output)
    generated_files = reporter.generate_all_reports()
    
    print(f"\n{'='*70}")
    print("✅ COMPLETE!")
    print(f"{'='*70}\n")
    
    # Summary
    print("Generated Reports:")
    print(f"\nPairwise Comparisons ({len(comparator.pairwise_comparisons)}):")
    for comp in comparator.pairwise_comparisons:
        print(f"  • {comp['schema1']} vs {comp['schema2']}: Excel + Word + HTML")
    
    print(f"\nMaster Reports:")
    print(f"  • {args.output}_MASTER_MATRIX.xlsx")
    print(f"  • {args.output}_MANAGEMENT_SUMMARY.docx")
    print(f"  • {args.output}_BUSINESS_IMPACT.docx")
    print(f"  • {args.output}_DEVELOPER_GUIDE.docx")
    print(f"  • {args.output}_QA_CHECKLIST.docx")
    
    print(f"\n📊 Total files generated: {len(generated_files)}")
    print()


if __name__ == '__main__':
    main()
